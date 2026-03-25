"""
Aplikasi Python untuk membaca dokumen (PDF, DOC, DOCX, JPG, JPEG, PNG)
dari direktori server dan menyimpan hasilnya ke database PostgreSQL.
"""

import os
import io
from datetime import datetime
from pathlib import Path
from typing import Optional, List, Dict, Any

# Database
import psycopg2
from psycopg2.extras import RealDictCursor

# PDF handling
import PyPDF2

# DOC/DOCX handling
import docx
from docx import Document

# Image handling and OCR
from PIL import Image
import pytesseract

# Untuk membaca metadata file
import mimetypes

# Import extractor untuk data spesifik
from data_extractor import DocumentDataExtractor


class DocumentReader:
    """Kelas untuk membaca berbagai jenis dokumen."""
    
    SUPPORTED_EXTENSIONS = {
        '.pdf': 'pdf',
        '.doc': 'doc',
        '.docx': 'docx',
        '.jpg': 'image',
        '.jpeg': 'image',
        '.png': 'image'
    }
    
    def __init__(self):
        pass
    
    def read_pdf(self, file_path: str) -> Dict[str, Any]:
        """Membaca file PDF dan mengekstrak teks."""
        result = {
            'content': '',
            'page_count': 0,
            'metadata': {}
        }
        
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                result['page_count'] = len(reader.pages)
                
                # Ekstrak metadata
                if reader.metadata:
                    result['metadata'] = {
                        'title': reader.metadata.get('/Title', ''),
                        'author': reader.metadata.get('/Author', ''),
                        'subject': reader.metadata.get('/Subject', ''),
                        'creator': reader.metadata.get('/Creator', ''),
                        'producer': reader.metadata.get('/Producer', ''),
                        'creation_date': str(reader.metadata.get('/CreationDate', '')),
                        'mod_date': str(reader.metadata.get('/ModDate', ''))
                    }
                
                # Ekstrak teks dari semua halaman
                text_content = []
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        text_content.append(text)
                
                result['content'] = '\n\n'.join(text_content)
                
        except Exception as e:
            result['error'] = str(e)
        
        return result
    
    def read_docx(self, file_path: str) -> Dict[str, Any]:
        """Membaca file DOCX dan mengekstrak teks."""
        result = {
            'content': '',
            'paragraph_count': 0,
            'metadata': {}
        }
        
        try:
            doc = Document(file_path)
            
            # Ekstrak metadata dari core properties
            if doc.core_properties:
                result['metadata'] = {
                    'title': doc.core_properties.title or '',
                    'author': doc.core_properties.author or '',
                    'subject': doc.core_properties.subject or '',
                    'keywords': doc.core_properties.keywords or '',
                    'created': str(doc.core_properties.created) if doc.core_properties.created else '',
                    'modified': str(doc.core_properties.modified) if doc.core_properties.modified else ''
                }
            
            # Ekstrak teks dari paragraf
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            result['paragraph_count'] = len(paragraphs)
            result['content'] = '\n\n'.join(paragraphs)
            
        except Exception as e:
            result['error'] = str(e)
        
        return result
    
    def read_doc(self, file_path: str) -> Dict[str, Any]:
        """Membaca file DOC (format lama)."""
        # Catatan: Format .doc lama lebih sulit dibaca tanpa library khusus
        # Kita coba baca sebagai docx atau gunakan konversi
        result = {
            'content': '',
            'metadata': {},
            'warning': 'File .doc format lama mungkin memerlukan konversi terlebih dahulu'
        }
        
        try:
            # Coba baca dengan docx (kadang berhasil jika format kompatibel)
            doc = Document(file_path)
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            result['content'] = '\n\n'.join(paragraphs)
            result['metadata']['note'] = 'Dibaca dengan library python-docx'
            
        except Exception as e:
            result['error'] = f'Tidak dapat membaca file .doc langsung: {str(e)}'
            result['content'] = ''
        
        return result
    
    def read_image(self, file_path: str, lang: str = 'eng+ind') -> Dict[str, Any]:
        """Membaca file gambar dan melakukan OCR untuk mengekstrak teks."""
        result = {
            'content': '',
            'metadata': {}
        }
        
        try:
            # Buka gambar
            img = Image.open(file_path)
            
            # Ekstrak metadata gambar
            result['metadata'] = {
                'format': img.format,
                'mode': img.mode,
                'size': img.size,
                'width': img.width,
                'height': img.height
            }
            
            # Ekstrak informasi EXIF jika ada
            if hasattr(img, '_getexif') and img._getexif():
                result['metadata']['exif'] = str(img._getexif())
            
            # Lakukan OCR untuk mengekstrak teks
            # Pastikan tesseract-ocr sudah terinstall di sistem
            text = pytesseract.image_to_string(img, lang=lang)
            result['content'] = text.strip()
            
        except Exception as e:
            result['error'] = str(e)
        
        return result
    
    def read_file(self, file_path: str) -> Dict[str, Any]:
        """Membaca file berdasarkan ekstensinya."""
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        
        if extension not in self.SUPPORTED_EXTENSIONS:
            return {
                'error': f'Ekstensi file {extension} tidak didukung',
                'content': '',
                'metadata': {}
            }
        
        file_type = self.SUPPORTED_EXTENSIONS[extension]
        
        # Baca metadata file sistem
        file_stats = os.stat(file_path)
        basic_metadata = {
            'filename': file_path.name,
            'filepath': str(file_path.absolute()),
            'size_bytes': file_stats.st_size,
            'created_time': datetime.fromtimestamp(file_stats.st_ctime).isoformat(),
            'modified_time': datetime.fromtimestamp(file_stats.st_mtime).isoformat(),
            'mime_type': mimetypes.guess_type(str(file_path))[0] or 'unknown'
        }
        
        # Baca konten berdasarkan tipe file
        if file_type == 'pdf':
            result = self.read_pdf(str(file_path))
        elif file_type == 'docx':
            result = self.read_docx(str(file_path))
        elif file_type == 'doc':
            result = self.read_doc(str(file_path))
        elif file_type == 'image':
            result = self.read_image(str(file_path))
        else:
            result = {'error': 'Tipe file tidak dikenali', 'content': '', 'metadata': {}}
        
        # Gabungkan metadata
        if 'metadata' in result:
            basic_metadata.update(result['metadata'])
        result['metadata'] = basic_metadata
        
        return result


class DocumentDatabase:
    """Kelas untuk mengelola penyimpanan dokumen ke PostgreSQL."""
    
    def __init__(self, connection_params: Dict[str, str]):
        """
        Inisialisasi koneksi database.
        
        Args:
            connection_params: Dictionary dengan parameter koneksi:
                - host: Host database
                - port: Port database (default: 5432)
                - database: Nama database
                - user: Username
                - password: Password
        """
        self.connection_params = connection_params
        self.connection = None
        self._connect()
    
    def _connect(self):
        """Membuat koneksi ke database PostgreSQL."""
        try:
            self.connection = psycopg2.connect(
                host=self.connection_params.get('host', 'localhost'),
                port=self.connection_params.get('port', '5432'),
                database=self.connection_params.get('database'),
                user=self.connection_params.get('user'),
                password=self.connection_params.get('password')
            )
            print("✓ Koneksi ke PostgreSQL berhasil")
        except Exception as e:
            print(f"✗ Gagal koneksi ke PostgreSQL: {e}")
            raise
    
    def create_tables(self):
        """Membuat tabel untuk menyimpan dokumen jika belum ada."""
        create_table_query = """
        CREATE TABLE IF NOT EXISTS documents (
            id SERIAL PRIMARY KEY,
            filename VARCHAR(255) NOT NULL,
            filepath TEXT NOT NULL,
            file_type VARCHAR(50),
            file_size_bytes BIGINT,
            mime_type VARCHAR(100),
            content TEXT,
            page_count INTEGER,
            paragraph_count INTEGER,
            metadata JSONB,
            extracted_data JSONB,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            processing_status VARCHAR(50) DEFAULT 'processed',
            error_message TEXT
        );
        
        CREATE INDEX IF NOT EXISTS idx_documents_filename ON documents(filename);
        CREATE INDEX IF NOT EXISTS idx_documents_filepath ON documents(filepath);
        CREATE INDEX IF NOT EXISTS idx_documents_file_type ON documents(file_type);
        CREATE INDEX IF NOT EXISTS idx_documents_created_at ON documents(created_at);
        CREATE INDEX IF NOT EXISTS idx_extracted_invoice ON documents USING GIN (extracted_data);
        """
        
        try:
            with self.connection.cursor() as cursor:
                cursor.execute(create_table_query)
                self.connection.commit()
                print("✓ Tabel documents berhasil dibuat/diverifikasi")
        except Exception as e:
            print(f"✗ Gagal membuat tabel: {e}")
            self.connection.rollback()
            raise
    
    def save_document(self, file_path: str, read_result: Dict[str, Any], 
                     extracted_data: Optional[Dict] = None) -> Optional[int]:
        """
        Menyimpan hasil pembacaan dokumen ke database.
        
        Args:
            file_path: Path lengkap file
            read_result: Hasil pembacaan dari DocumentReader
            extracted_data: Data yang diekstrak (invoice number, tax ID, dll)
            
        Returns:
            ID dokumen yang disimpan, atau None jika gagal
        """
        insert_query = """
        INSERT INTO documents (
            filename, filepath, file_type, file_size_bytes, mime_type,
            content, page_count, paragraph_count, metadata, extracted_data,
            processing_status, error_message
        ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
        RETURNING id;
        """
        
        try:
            metadata = read_result.get('metadata', {})
            error = read_result.get('error')
            
            # Tentukan tipe file dari ekstensi
            file_ext = Path(file_path).suffix.lower()
            file_type_map = {
                '.pdf': 'pdf',
                '.doc': 'doc',
                '.docx': 'docx',
                '.jpg': 'image',
                '.jpeg': 'image',
                '.png': 'image'
            }
            file_type = file_type_map.get(file_ext, 'unknown')
            
            values = (
                metadata.get('filename', Path(file_path).name),
                metadata.get('filepath', str(Path(file_path).absolute())),
                file_type,
                metadata.get('size_bytes'),
                metadata.get('mime_type'),
                read_result.get('content', ''),
                read_result.get('page_count'),
                read_result.get('paragraph_count'),
                psycopg2.extras.Json(metadata),
                psycopg2.extras.Json(extracted_data or {}),
                'error' if error else 'processed',
                error
            )
            
            with self.connection.cursor() as cursor:
                cursor.execute(insert_query, values)
                doc_id = cursor.fetchone()[0]
                self.connection.commit()
                print(f"✓ Dokumen tersimpan dengan ID: {doc_id}")
                return doc_id
                
        except Exception as e:
            print(f"✗ Gagal menyimpan dokumen: {e}")
            self.connection.rollback()
            return None
    
    def get_all_documents(self, limit: int = 100) -> List[Dict]:
        """Mengambil semua dokumen dari database."""
        query = "SELECT * FROM documents ORDER BY created_at DESC LIMIT %s;"
        
        try:
            with self.connection.cursor(cursor_factory=RealDictCursor) as cursor:
                cursor.execute(query, (limit,))
                return cursor.fetchall()
        except Exception as e:
            print(f"✗ Gagal mengambil dokumen: {e}")
            return []
    
    def search_documents(self, keyword: str, limit: int = 100) -> List[Dict]:
        """Mencari dokumen berdasarkan kata kunci dalam konten."""
        query = """
        SELECT * FROM documents 
        WHERE content ILIKE %s OR filename ILIKE %s
        ORDER BY created_at DESC 
        LIMIT %s;
        """
        
        try:
            search_pattern = f"%{keyword}%"
            with self.connection.cursor(cursor_factory=RealDictCursor) as cursor:
                cursor.execute(query, (search_pattern, search_pattern, limit))
                return cursor.fetchall()
        except Exception as e:
            print(f"✗ Gagal mencari dokumen: {e}")
            return []

    def search_by_extracted_data(self, field: str, value: str, limit: int = 100) -> List[Dict]:
        """
        Mencari dokumen berdasarkan data yang diekstrak (invoice number, NPWP, dll).
        
        Args:
            field: Field yang dicari (invoice_number, tax_invoice_number, npwp, date, amount)
            value: Nilai yang dicari
            limit: Batas hasil
            
        Returns:
            List dokumen yang match
        """
        query = """
        SELECT * FROM documents 
        WHERE extracted_data->%s @> %s::jsonb
        ORDER BY created_at DESC 
        LIMIT %s;
        """
        
        try:
            # Format value sebagai JSON array untuk pencarian
            import json
            json_value = json.dumps([value])
            
            with self.connection.cursor(cursor_factory=RealDictCursor) as cursor:
                cursor.execute(query, (field, json_value, limit))
                return cursor.fetchall()
        except Exception as e:
            print(f"✗ Gagal mencari dokumen berdasarkan extracted data: {e}")
            return []

    def get_documents_by_date_range(self, start_date: str, end_date: str, limit: int = 100) -> List[Dict]:
        """
        Mendapatkan dokumen dalam rentang tanggal tertentu.
        
        Args:
            start_date: Tanggal mulai (YYYY-MM-DD)
            end_date: Tanggal akhir (YYYY-MM-DD)
            limit: Batas hasil
        """
        query = """
        SELECT * FROM documents 
        WHERE extracted_data->'date' IS NOT NULL
        AND EXISTS (
            SELECT 1 FROM jsonb_array_elements_text(extracted_data->'date') AS date_val
            WHERE date_val BETWEEN %s AND %s
        )
        ORDER BY created_at DESC 
        LIMIT %s;
        """
        
        try:
            with self.connection.cursor(cursor_factory=RealDictCursor) as cursor:
                cursor.execute(query, (start_date, end_date, limit))
                return cursor.fetchall()
        except Exception as e:
            print(f"✗ Gagal mencari dokumen berdasarkan rentang tanggal: {e}")
            return []
    
    def close(self):
        """Menutup koneksi database."""
        if self.connection:
            self.connection.close()
            print("✓ Koneksi database ditutup")


class DocumentProcessor:
    """Kelas utama untuk memproses dokumen dari direktori."""
    
    def __init__(self, db_config: Dict[str, str]):
        """
        Inisialisasi processor dokumen.
        
        Args:
            db_config: Konfigurasi koneksi database PostgreSQL
        """
        self.reader = DocumentReader()
        self.db = DocumentDatabase(db_config)
        self.stats = {
            'total_files': 0,
            'processed': 0,
            'failed': 0,
            'skipped': 0
        }
    
        # Inisialisasi extractor untuk data spesifik
        self.extractor = DocumentDataExtractor()
    
    def process_directory(self, directory_path: str, recursive: bool = True, 
                         extract_data: bool = True) -> Dict[str, int]:
        """
        Memproses semua dokumen dalam direktori.
        
        Args:
            directory_path: Path ke direktori yang akan diproses
            recursive: Jika True, proses juga subdirektori
            
        Returns:
            Statistik pemrosesan
        """
        directory = Path(directory_path)
        
        if not directory.exists():
            print(f"✗ Direktori tidak ditemukan: {directory_path}")
            return self.stats
        
        if not directory.is_dir():
            print(f"✗ Path bukan direktori: {directory_path}")
            return self.stats
        
        print(f"\n📁 Memproses direktori: {directory.absolute()}")
        print("=" * 60)
        
        # Buat tabel jika belum ada
        self.db.create_tables()
        
        # Dapatkan semua file yang didukung
        files_to_process = []
        if recursive:
            for ext in DocumentReader.SUPPORTED_EXTENSIONS.keys():
                files_to_process.extend(directory.rglob(f"*{ext}"))
                files_to_process.extend(directory.rglob(f"*{ext.upper()}"))
        else:
            for ext in DocumentReader.SUPPORTED_EXTENSIONS.keys():
                files_to_process.extend(directory.glob(f"*{ext}"))
                files_to_process.extend(directory.glob(f"*{ext.upper()}"))
        
        # Hapus duplikat
        files_to_process = list(set(files_to_process))
        self.stats['total_files'] = len(files_to_process)
        
        print(f"Ditemukan {len(files_to_process)} file untuk diproses\n")
        
        # Proses setiap file
        for file_path in files_to_process:
            self.stats['total_files'] += 1
            print(f"📄 Memproses: {file_path.name}")
            
            try:
                # Baca file
                result = self.reader.read_file(str(file_path))
                
                if 'error' in result and result.get('content') == '':
                    print(f"   ⚠️  Peringatan: {result.get('error', 'Unknown error')}")
                    self.stats['failed'] += 1
                else:
                    # Ekstrak data spesifik jika diminta
                    extracted_data = None
                    if extract_data and result.get('content'):
                        extracted_data = self.extractor.extract_all(result['content'])
                        if extracted_data:
                            print(f"   📊 Data diekstrak: {list(extracted_data.keys())}")
                    
                    # Simpan ke database
                    doc_id = self.db.save_document(str(file_path), result, extracted_data)
                    if doc_id:
                        content_preview = result.get('content', '')[:100]
                        if content_preview:
                            print(f"   ✓ Berhasil ({len(content_preview)} chars)")
                        else:
                            print(f"   ✓ Berhasil (no text content)")
                        self.stats['processed'] += 1
                    else:
                        self.stats['failed'] += 1
                        
            except Exception as e:
                print(f"   ✗ Error: {str(e)}")
                self.stats['failed'] += 1
        
        print("\n" + "=" * 60)
        print(f"✅ Pemrosesan selesai!")
        print(f"   Total file: {self.stats['total_files']}")
        print(f"   Berhasil: {self.stats['processed']}")
        print(f"   Gagal: {self.stats['failed']}")
        print("=" * 60)
        
        return self.stats
    
    def process_single_file(self, file_path: str) -> Optional[int]:
        """
        Memproses satu file dokumen.
        
        Args:
            file_path: Path ke file yang akan diproses
            
        Returns:
            ID dokumen yang disimpan, atau None jika gagal
        """
        file = Path(file_path)
        
        if not file.exists():
            print(f"✗ File tidak ditemukan: {file_path}")
            return None
        
        print(f"📄 Memproses: {file.name}")
        
        # Buat tabel jika belum ada
        self.db.create_tables()
        
        try:
            result = self.reader.read_file(str(file))
            doc_id = self.db.save_document(str(file), result)
            
            if doc_id:
                print(f"   ✓ Berhasil disimpan dengan ID: {doc_id}")
            else:
                print(f"   ✗ Gagal menyimpan")
            
            return doc_id
            
        except Exception as e:
            print(f"   ✗ Error: {str(e)}")
            return None
    
    def close(self):
        """Menutup semua resource."""
        self.db.close()


def main():
    """Fungsi utama untuk menjalankan aplikasi."""
    
    # Konfigurasi database PostgreSQL
    DB_CONFIG = {
        'host': 'localhost',
        'port': '5432',
        'database': 'document_db',
        'user': 'postgres',
        'password': 'your_password_here'  # Ganti dengan password Anda
    }
    
    # Direktori yang berisi dokumen
    DOCUMENT_DIRECTORY = '/path/to/your/documents'  # Ganti dengan path direktori Anda
    
    print("=" * 60)
    print("📚 Aplikasi Pembaca Dokumen")
    print("=" * 60)
    print("\nFormat yang didukung:")
    print("  - PDF (.pdf)")
    print("  - Word (.doc, .docx)")
    print("  - Gambar (.jpg, .jpeg, .png)")
    print("\nFitur Ekstraksi Data:")
    print("  - Nomor Invoice")
    print("  - Nomor Faktur Pajak")
    print("  - NPWP")
    print("  - Tanggal")
    print("  - Jumlah Uang (Amount)")
    print("  - Nama Perusahaan")
    print("\n" + "=" * 60)
    
    # Inisialisasi processor
    try:
        processor = DocumentProcessor(DB_CONFIG)
        
        # Proses direktori dengan ekstraksi data
        stats = processor.process_directory(DOCUMENT_DIRECTORY, recursive=True, extract_data=True)
        
        # Tampilkan contoh dokumen yang tersimpan
        print("\n📋 Contoh dokumen tersimpan:")
        docs = processor.db.get_all_documents(limit=5)
        for doc in docs:
            print(f"  - {doc['filename']} (ID: {doc['id']}, Type: {doc['file_type']})")
            # Tampilkan data yang diekstrak jika ada
            if doc.get('extracted_data'):
                ext = doc['extracted_data']
                if ext.get('invoice_number'):
                    print(f"      Invoice: {ext['invoice_number']}")
                if ext.get('tax_invoice_number'):
                    print(f"      Faktur Pajak: {ext['tax_invoice_number']}")
                if ext.get('date'):
                    print(f"      Tanggal: {ext['date']}")
        
        # Contoh pencarian berdasarkan data ekstrak
        print("\n\n🔍 Contoh Pencarian:")
        print("-" * 40)
        
        # Cari berdasarkan nomor invoice
        print("\nCari Invoice INV-2023-998:")
        results = processor.db.search_by_extracted_data('invoice_number', 'INV-2023-998')
        for doc in results:
            print(f"  Ditemukan: {doc['filename']}")
        
        # Cari berdasarkan rentang tanggal
        print("\nCari dokumen tanggal 01/01/2023 - 31/12/2023:")
        results = processor.db.get_documents_by_date_range('2023-01-01', '2023-12-31')
        for doc in results:
            print(f"  Ditemukan: {doc['filename']} - {doc.get('extracted_data', {}).get('date', [])}")
        
        processor.close()
        
    except Exception as e:
        print(f"\n✗ Error: {e}")
        print("\nPastikan:")
        print("  1. PostgreSQL sudah running")
        print("  2. Database sudah dibuat")
        print("  3. Kredensial database benar")
        print("  4. Direktori dokumen ada dan dapat diakses")


if __name__ == "__main__":
    main()
